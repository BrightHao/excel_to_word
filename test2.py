import xlrd
from docx import Document
from docx.shared import Pt


workbook = xlrd.open_workbook('C:\\Users\\ASUS\\Desktop\\2014届电气学院_选题表汇总提交1219.xls')
sheet_name = workbook.sheet_names()[6]
#实例化
sheet = workbook.sheet_by_name(sheet_name)
# 获取行数
print(sheet_name, sheet.nrows)
#创建一个新的word文档
document = Document()
p = document.add_paragraph()

for i in range(1, sheet.nrows):
    name = sheet.row_values(i)[4]
    Class = sheet.row_values(i)[2]
    teacher = sheet.row_values(i)[8]
    content_head = "2019届毕业设计指导记录本"
    text_head = p.add_run(content_head)
    text_head.font.name = '宋体'
    text_head.font.size = Pt(15)
    text_head.bold = True
    text = p.add_run(f"""
学    院：电气与信息工程学院
学生姓名：{name}
班级名称：{Class}
指导老师：{teacher}""")
    text.font.name = "宋体"
    text.font.size = Pt(14)
    # block = p.add_run("""
#
#
#
# """)
#     block.font.size = Pt(12)
    if i % 4 != 0:
        block = p.add_run("""




""")
        block.font.size = Pt(12)
    else:
        p.add_run("""
""")

document.save('C:\\Users\\ASUS\\Desktop\\test5.docx')
#a = p.add_run('bold')
#a.bold = True
#a.font.size = "13.5"


