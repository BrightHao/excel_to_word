import xlrd
from docx import Document
from docx.shared import Pt


workbook = xlrd.open_workbook('C:\\Users\\Tang Haolun\\Desktop\\2021届电气学院毕业设计选题一览表完整版.xls')
sheet_name = "选题表"
#实例化
sheet = workbook.sheet_by_name(sheet_name)
# 获取行数
print(sheet_name, sheet.nrows)
#创建一个新的word文档
document = Document()
p = document.add_paragraph()

for i in range(1, sheet.nrows):
    name = sheet.row_values(i)[4]
    Class = sheet.row_values(i)[5]
    teacher = sheet.row_values(i)[7]
    if i == 1:
        content = f"""


   2021届毕业设计指导记录本 """
    elif i % 4 == 1 and i > 1:
        content = u"""





   2021届毕业设计指导记录本 """
    else:
        content = u"""



   2021届毕业设计指导记录本 """
    text_head = p.add_run(content)
    text_head.font.name = '宋体'
    text_head.font.size = Pt(13.5)
    text_head.bold = True
    text = p.add_run(f"""
学    院：电气与信息工程学院
学生姓名：{name}
班级名称：{Class}
指导老师：{teacher}""")
    text.font.name = "宋体"
    text.font.size = Pt(12)
document.save('C:\\Users\\Tang Haolun\\Desktop\\2021届电气学院毕业设计选题封面.docx')
#a = p.add_run('bold')
#a.bold = True
#a.font.size = "13.5"


# """
# 学    院：电气与信息工程学院
# 学生姓名：{name}
# 班级名称：{Class}
# 指导老师：{teacher}
# """
