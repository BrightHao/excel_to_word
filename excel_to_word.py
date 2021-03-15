import xlrd
from docx import Document
from docx.shared import Pt


workbook = xlrd.open_workbook('C:\\Users\\Tang Haolun\\Desktop\\2021届电气学院毕业设计选题一览表完整版.xls') # 被读取execel文件的路径
sheet_name = "选题表"  # 表格名称
#实例化
sheet = workbook.sheet_by_name(sheet_name)
# 获取行数
print(sheet_name, sheet.nrows)
#创建一个新的word文档
document = Document()
p = document.add_paragraph()

for i in range(1, sheet.nrows):
    name = sheet.row_values(i)[4]
    Class = sheet.row_values(i)[5]
    teacher = sheet.row_values(i)[7]
    content_head = "2021届毕业设计指导记录本"
    text_head = p.add_run(content_head)
    text_head.font.name = '宋体'
    text_head.font.size = Pt(15)
    text_head.bold = True
    text = p.add_run(f"""
学    院：湖南工业大学电气学院
学生姓名：{name}
班级名称：{Class}
指导老师：{teacher}
""")
    text.font.name = "宋体"
    text.font.size = Pt(14)
    if i % 4 == 0:##删除五格
        block = p.add_run("""
""")
        block.font.size = Pt(9)
    else:
        block = p.add_run("""





""")
        block.font.size = Pt(9)
document.save('C:\\Users\\Tang Haolun\\Desktop\\毕设选题封面.docx')  # 保存的文件路径和名称

