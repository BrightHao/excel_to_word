# -*- coding: utf-8 -*-

# document_1.write('F:\\{}.docx'.format(emp_full_info.get("identity_card")))
# Year=year.decode(encoding='UTF-8', errors='strict')

# from mailmerge import MailMerge
# import xlrd
#
# workbook = xlrd.open_workbook('C:\\Users\\ASUS\\Desktop\\2014届电气学院_选题表汇总提交1219.xls')
# sheet_name = workbook.sheet_names()[6]
# #实例化
# sheet = workbook.sheet_by_name(sheet_name)
# # 获取行数
# print(sheet_name, sheet.nrows)
#
#
# # 打印模板
# template = "C:\\Users\\ASUS\\Desktop\\2019届毕业设计指导记录本.docx"
# # 创建邮件合并文档并查看所有字段
#
#
# # print("Fields included in {}: {}".format(template, document_1.get_merge_fields()))
# list = []
# for i in range(1, sheet.nrows):
#     name = sheet.row_values(i)[4]
#     Class = sheet.row_values(i)[2]
#     teacher = sheet.row_values(i)[8]
#     # print(f'{i}, {name}, {Class}, {teacher}')
#     list.append([name, Class, teacher])
#     if i % 8 == 0:
#         document_1 = MailMerge(template)
#         document_1.merge(
#             name=u'{}'.format(list[0][0]),
#             Class=u'{}'.format(list[0][1]),
#             teacher=u'{}'.format(list[0][2]),
#             name1=u'{}'.format(list[1][0]),
#             Class1=u'{}'.format(list[1][1]),
#             teacher1=u'{}'.format(list[1][2]),
#             name2=u'{}'.format(list[2][0]),
#             Class2=u'{}'.format(list[2][1]),
#             teacher2=u'{}'.format(list[2][2]),
#             name3=u'{}'.format(list[3][0]),
#             Class3=u'{}'.format(list[3][1]),
#             teacher3=u'{}'.format(list[3][2]),
#             name4=u'{}'.format(list[4][0]),
#             Class4=u'{}'.format(list[4][1]),
#             teacher4=u'{}'.format(list[4][2]),
#             name5=u'{}'.format(list[5][0]),
#             Class5=u'{}'.format(list[5][1]),
#             teacher5=u'{}'.format(list[5][2]),
#             name6=u'{}'.format(list[6][0]),
#             Class6=u'{}'.format(list[6][1]),
#             teacher6=u'{}'.format(list[6][2]),
#             name7=u'{}'.format(list[7][0]),
#             Class7=u'{}'.format(list[7][1]),
#             teacher7=u'{}'.format(list[7][2]),
#         )
#         document_1.write('C:\\Users\\ASUS\\Desktop\\新建文件夹\\{}.docx'.format(list[0][0]))
#         list = []

from docx import Document
from docx.shared import Pt

dir = "C:\\Users\\ASUS\\Desktop\\2019届毕业设计指导记录本.docx"
document = Document(dir)
p = document.add_paragraph()
#字体加粗
# p.add_run('bold').bold = True
#斜字体
# p.add_run('italic.').italic = True
text = p.add_run('dfsgklkjhfeswdqaSWDEFGHbHKSBFZIUJN金佛的深加工及加热都给解耦我热手工我结算日偶家 ')
font = text.font
font.name = u'宋体'
font.size = Pt(10.5)


document.save('C:\\Users\\ASUS\\Desktop\\test.docx')


# document = Document()
# paragraph = document.add_paragraph()
#
# run = paragraph.add_run("hellohellohello")
# font = run.font
#
# 设置字体样式
# font.name = u'宋体'
# 设置字体大小
# font.size = Pt(55)
document.save('C:\\Users\\ASUS\\Desktop\\test.docx')
