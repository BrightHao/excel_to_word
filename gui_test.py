#!/usr/bin/env python
# -*- coding: utf-8 -*-

from tkinter import *
from tkinter import filedialog
from tkinter import messagebox

import xlrd
from docx import Document
from docx.shared import Pt



class Windows():
    def __init__(self, window):
        self.win = window

    #设置窗口
    def set_init_window(self):
        self.win.title("毕设封面导出工具v1.0  QQ:861759757")  # 窗口名
        self.win.geometry("550x450")

        # 输入框
        # 选择文件
        excel_path = Variable()
        self.c_excel_path = Entry(self.win, text=excel_path, width=50, highlightcolor='red', highlightthickness=1)
        self.c_excel_path.place(x=140, y=-40+100)

        # 保存文件名
        save_name = Variable()
        self.c_save_name = Entry(self.win, text=save_name, width=50, highlightcolor='red', highlightthickness=1)
        self.c_save_name.place(x=140, y=0+100)

        sheet = Variable()
        self.c_sheet = Entry(self.win, text=sheet, width=50, highlightcolor='red', highlightthickness=1)
        self.c_sheet.place(x=140, y=40+100)

        student = Variable()
        self.c_student = Entry(self.win, text=student, width=50, highlightcolor='red', highlightthickness=1)
        self.c_student.place(x=140, y=80+100)

        clas = Variable()
        self.c_clas = Entry(self.win, text=clas, width=50, highlightcolor='red', highlightthickness=1)
        self.c_clas.place(x=140, y=120+100)

        teacher = Variable()
        self.c_teacher = Entry(self.win, text=teacher, width=50, highlightcolor='red', highlightthickness=1)
        self.c_teacher.place(x=140, y=160+100)

        # 选择文件按钮
        self.chooseFileBtn = Button(self.win, text='选择文件', command=self.choose_file).place(x=75, y=-40+100)
        # 保存文件按钮
        self.saveFileBtn = Button(self.win, text='输出目录', command=self.save_file).place(x=75, y=0+100)

        # 标签
        Label(self.win, text='表格名').place(x=75, y=40+100)
        Label(self.win, text='学生列').place(x=75, y=80+100)
        Label(self.win, text='班级列').place(x=75, y=120+100)
        Label(self.win, text='指导老师列').place(x=75, y=160+100)

        # 开始按钮
        Button(self.win, text="开始转换", bg="lightblue", command=self.translate).place(x=230, y=320)

    # 选择文件
    def choose_file(self):
        selectFile = filedialog.askopenfilename(filetypes=[("XLS", ".xls"), ("XLSX", ".xlsx")])
        self.c_excel_path.delete(0, "end")
        self.c_excel_path.insert(0, selectFile)

    # 保存文件
    def save_file(self):
        saveFile = filedialog.asksaveasfilename(filetypes=[("docx", ".docx")])
        if len(saveFile) > 0 and saveFile[-5:] != ".docx":
            saveFile += ".docx"
        self.c_save_name.delete(0, "end")
        self.c_save_name.insert(0, saveFile)

    def translate(self):
        excel_name = self.c_excel_path.get()
        save_name = self.c_save_name.get()
        sheet_name = self.c_sheet.get()
        clas = self.c_clas.get()
        student = self.c_student.get()
        teacher = self.c_teacher.get()

        workbook = xlrd.open_workbook(excel_name)
        #实例化
        sheet = workbook.sheet_by_name(sheet_name)
        #创建一个新的word文档
        document = Document()
        p = document.add_paragraph()

        class_index = ord(clas)-65
        student_index = ord(student)-65
        teacher_index = ord(teacher)-65
        for i in range(1, sheet.nrows):
            name = sheet.row_values(i)[student_index]
            Class = sheet.row_values(i)[class_index]
            teacher = sheet.row_values(i)[teacher_index]
            content_head = "2021届毕业设计指导记录本"
            text_head = p.add_run(content_head)
            text_head.font.name = '宋体'
            text_head.font.size = Pt(15)
            text_head.bold = True
            text = p.add_run(f"""
学    院：电气与信息工程学院
学生姓名：{name}
班级名称：{Class}
指导老师：{teacher}
""")
            text.font.name = "宋体"
            text.font.size = Pt(14)
            if i % 4 == 0:##删除五格
                block = p.add_run("""
""")
                block.font.size = Pt(9)
            else:
                block = p.add_run("""





""")
                block.font.size = Pt(9)
        document.save(save_name)  # 保存的文件路径和名称
        messagebox.showinfo('提示','转换完毕!')


def gui_start():
    init_window = Tk()              #实例化出一个父窗口
    ZMJ_PORTAL = Windows(init_window)
    # 设置根窗口默认属性
    ZMJ_PORTAL.set_init_window()

    init_window.mainloop()          #父窗口进入事件循环，可以理解为保持窗口运行，否则界面不展示


if __name__ == "__main__":
    gui_start()

